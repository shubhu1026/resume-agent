from docx import Document

def is_section_header(para):
    return para.style.name == "Normal" and para.text.isupper() and para.text.strip() != ""

def replace_section_flexible(doc, section_title, new_content_lines):
    paragraphs = doc.paragraphs
    inside_section = False
    section_paragraphs = []

    # Find the section paragraphs between headers
    for para in paragraphs:
        if para.text.strip().upper() == section_title.upper():
            inside_section = True
            continue
        if inside_section:
            if is_section_header(para):
                break
            section_paragraphs.append(para)

    if not section_paragraphs:
        return f"❌ Section '{section_title}' not found."

    # Update existing paragraphs or add new ones
    for i in range(min(len(section_paragraphs), len(new_content_lines))):
        para = section_paragraphs[i]
        para.clear()  # clear existing text

        line = new_content_lines[i]

        if section_title.upper() == "SKILLS":
            line = line.lstrip("• ").strip()
            if ':' in line:
                category, rest = line.split(':', 1)
                para.add_run("• ").bold = False
                run = para.add_run(f"{category.strip()}: ")
                run.bold = True
                para.add_run(rest.strip())
            else:
                para.add_run("• " + line)

        elif section_title.upper() == "PROJECTS":
            if line.startswith('•'):
                para.add_run(line).bold = False
            else:
                para.add_run(line).bold = True

        else:  # SUMMARY and other sections
            para.text = line

    # Remove any extra old paragraphs
    if len(new_content_lines) < len(section_paragraphs):
        for para in section_paragraphs[len(new_content_lines):]:
            p = para._element
            p.getparent().remove(p)

    # Append any extra new lines if needed
    elif len(new_content_lines) > len(section_paragraphs):
        last_para = section_paragraphs[-1]
        for line in new_content_lines[len(section_paragraphs):]:
            p = last_para._parent.add_paragraph()
            if section_title.upper() == "SKILLS":
                line = line.lstrip("• ").strip()
                if ':' in line:
                    category, rest = line.split(':', 1)
                    p.add_run("• ").bold = False
                    run = p.add_run(f"{category.strip()}: ")
                    run.bold = True
                    p.add_run(rest.strip())
                else:
                    p.add_run("• " + line)
            elif section_title.upper() == "PROJECTS":
                if line.startswith('•'):
                    p.add_run(line).bold = False
                else:
                    p.add_run(line).bold = True
            else:
                p.text = line
            p.style = last_para.style

    return f"✅ Section '{section_title}' updated successfully."

def clean_bullets(lines):
    return [line.lstrip("• ").strip() for line in lines.splitlines() if line.strip()]

def clean_bullets_in_doc(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            # Fix patterns like '• • ' → '• '
            if "• •" in run.text:
                run.text = run.text.replace("• •", "• ")
            # Also fix leading spaces like '  • '
            if run.text.lstrip().startswith("•  "):
                run.text = run.text.lstrip().replace("•  ", "• ")

def update_resume(summary_lines, skills_lines, projects_lines, resume_path='your_resume.docx', output_path='updated_resume.docx'):
    doc = Document(resume_path)

    if isinstance(summary_lines, str):
        summary_lines = [line.strip() for line in summary_lines.split('\n') if line.strip()]
    if isinstance(skills_lines, str):
        skills_lines = [line.strip() for line in skills_lines.split('\n') if line.strip()]
    if isinstance(projects_lines, str):
        projects_lines = [line.strip() for line in projects_lines.split('\n') if line.strip()]

    msg1 = replace_section_flexible(doc, "SUMMARY", summary_lines)
    msg2 = replace_section_flexible(doc, "SKILLS", skills_lines)
    msg3 = replace_section_flexible(doc, "PROJECTS", projects_lines)

    # 🚨 Important: Cleanup redundant bullets from all runs
    clean_bullets_in_doc(doc)

    doc.save(output_path)
    return msg1, msg2, msg3, output_path
